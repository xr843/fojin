"""Parse uploaded chat attachments to plain text.

Supports PDF / TXT / MD / DOCX / CSV / HTML(.htm). The parsed text is
later prepended to the user's chat message, so we cap output at
``MAX_PARSED_CHARS`` (~20k tokens) to keep one big PDF from blowing past
the LLM context window.

Design notes:
- ``parse_attachment`` raises ``ValueError`` for unsupported extensions
  and corrupt/structurally-broken files. The chat upload API turns
  the ValueError into 422 with the message in the ``parse_error`` column,
  so the user gets a Chinese-readable hint instead of an opaque 500.
- An empty parse result is NOT an error; we return a placeholder string
  so the downstream "(无可提取文本)" hint is visible to the LLM and the
  user, instead of silently dropping the attachment.
- BeautifulSoup is preferred for HTML; if not installed we fall back to
  a regex tag-stripper. Both implementations are safe to run on
  untrusted input — we never execute scripts or fetch external
  resources.
"""

from __future__ import annotations

import csv
import io
import re

MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MiB hard cap on raw upload
MAX_PARSED_CHARS = 80_000  # ~20k tokens after parse; truncation marker appended
MAX_CSV_ROWS = 200  # render at most this many rows from a CSV
TRUNCATION_MARKER = "\n\n...（已截断）"
EMPTY_PLACEHOLDER = "(无可提取文本)"

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".html", ".htm"}


def _ext(filename: str) -> str:
    idx = filename.rfind(".")
    if idx == -1:
        return ""
    return filename[idx:].lower()


def _truncate(text: str) -> str:
    if len(text) <= MAX_PARSED_CHARS:
        return text
    return text[:MAX_PARSED_CHARS] + TRUNCATION_MARKER


def _parse_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def _parse_csv(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(decoded))
    rows: list[str] = []
    for i, row in enumerate(reader):
        if i >= MAX_CSV_ROWS:
            rows.append(f"...（仅显示前 {MAX_CSV_ROWS} 行）")
            break
        # Compact pipe-delimited render — easier for the LLM than raw CSV
        # because commas inside cell values won't confuse it.
        rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


_TAG_RE = re.compile(r"<[^>]+>")
_SCRIPT_STYLE_RE = re.compile(
    r"<(script|style)\b[^>]*>.*?</\1>", re.IGNORECASE | re.DOTALL
)
_WS_COLLAPSE_RE = re.compile(r"\n\s*\n\s*\n+")


def _parse_html(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup  # type: ignore[import-not-found]

        soup = BeautifulSoup(decoded, "html.parser")
        # Drop scripts/styles before extracting text — get_text would
        # otherwise dump JS/CSS source into the LLM prompt.
        for el in soup(["script", "style"]):
            el.decompose()
        text = soup.get_text(separator="\n", strip=True)
    except ImportError:
        # Fallback regex stripper. Less accurate (won't decode entities)
        # but never raises; good enough as a safety net if bs4 ever gets
        # uninstalled.
        stripped = _SCRIPT_STYLE_RE.sub("", decoded)
        text = _TAG_RE.sub("", stripped)
    return _WS_COLLAPSE_RE.sub("\n\n", text).strip()


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover — dep is required
        raise ValueError("PDF 解析依赖未安装") from exc

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"PDF 文件损坏或格式不支持：{exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:  # pragma: no cover — pypdf raises ad-hoc errors
            page_text = ""
        if page_text:
            pages.append(page_text)
    # Form-feed delimiter so the LLM can see page boundaries without us
    # injecting a noisy "===== Page N =====" header.
    return "\n\f\n".join(pages)


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover — dep is required
        raise ValueError("DOCX 解析依赖未安装") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"DOCX 文件损坏或格式不支持：{exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_attachment(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Parse an uploaded file to plain text.

    Returns the parsed text, truncated to ``MAX_PARSED_CHARS`` with a
    marker appended. Empty results return the ``EMPTY_PLACEHOLDER``
    string (not an exception) so the upload still succeeds and the user
    can see in the UI that the file was attempted but nothing useful
    was extracted.

    Raises ``ValueError`` for unsupported extensions or files that
    structurally fail to parse.
    """
    ext = _ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{ext or '(无扩展名)'}")

    if ext in (".txt", ".md"):
        text = _parse_text(file_bytes)
    elif ext == ".csv":
        text = _parse_csv(file_bytes)
    elif ext in (".html", ".htm"):
        text = _parse_html(file_bytes)
    elif ext == ".pdf":
        text = _parse_pdf(file_bytes)
    elif ext == ".docx":
        text = _parse_docx(file_bytes)
    else:  # pragma: no cover — set membership above guarantees coverage
        raise ValueError(f"不支持的文件类型：{ext}")

    text = text.strip()
    if not text:
        return EMPTY_PLACEHOLDER
    return _truncate(text)
