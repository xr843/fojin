"""Single-text document export — render one sutra (whole or one juan) as a
downloadable TXT / HTML / DOCX / EPUB document.

Design notes:
- The export is assembled from the stored per-juan plain content (already
  CBETA-cleaned, one source line per physical line). We split on newlines and
  drop blank lines; no reflow — a document export favours fidelity to the
  source lines over the reader's prose/verse heuristics.
- EPUB is built by hand with the stdlib ``zipfile`` (an EPUB is just a ZIP of
  XHTML + OPF + NCX). This keeps the dependency surface flat — no ebooklib.
- DOCX uses ``python-docx`` which is already a dependency.
- The pure ``render_*`` functions take an in-memory ``ExportDoc`` and return
  bytes, so they are unit-testable without a database.
"""

from dataclasses import dataclass, field
from html import escape
from io import BytesIO
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile, ZipInfo

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.text import BuddhistText, TextContent
from app.services.text_reflow import Segment, reflow

# format → (file extension, media type)
FORMATS: dict[str, tuple[str, str]] = {
    "txt": (".txt", "text/plain; charset=utf-8"),
    "html": (".html", "text/html; charset=utf-8"),
    "docx": (".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "epub": (".epub", "application/epub+zip"),
}

# EPUB 3 requires exactly one dcterms:modified. A fixed constant keeps re-exports
# byte-stable (the content itself, not the export time, is what identifies an
# edition) while satisfying epubcheck RSC-005.
_EPUB_MODIFIED = "2026-01-01T00:00:00Z"


@dataclass
class ExportJuan:
    juan_num: int
    # Raw stripped source lines — used by TXT (no layout) and as the reflow input.
    lines: list[str] = field(default_factory=list)
    # Reflowed layout segments — used by HTML/EPUB/DOCX to match the web reader.
    segments: list[Segment] = field(default_factory=list)


@dataclass
class ExportDoc:
    title: str
    cbeta_id: str
    translator: str | None = None
    dynasty: str | None = None
    juans: list[ExportJuan] = field(default_factory=list)


def _byline(doc: ExportDoc) -> str:
    """'後秦 佛陀耶舍譯'-style attribution line, omitting missing parts."""
    return " ".join(p for p in (doc.dynasty, doc.translator) if p)


def _juan_label(juan_num: int) -> str:
    return f"卷第{juan_num}"


# --- TXT ---------------------------------------------------------------------

def render_txt(doc: ExportDoc) -> str:
    out: list[str] = [doc.title]
    byline = _byline(doc)
    if byline:
        out.append(byline)
    out.append(f"CBETA {doc.cbeta_id}")
    for j in doc.juans:
        out.append("")
        out.append(_juan_label(j.juan_num))
        out.append("")
        out.extend(j.lines)
    return "\n".join(out) + "\n"


# --- HTML --------------------------------------------------------------------

# Mirrors the web reader's reader.css typography (.text-prose/.text-verse/…) so
# the exported page reads like fojin.app: flowing prose paragraphs (2em indent),
# indented verse, centred juan/section headings, right-aligned byline.
_HTML_STYLE = (
    "body{font-family:'Noto Serif SC','Source Han Serif SC',serif;max-width:42em;"
    "margin:2em auto;padding:0 1.2em;line-height:1.8;color:#1a1a1a;"
    "line-break:strict}"
    "h1{font-size:1.5em;margin:0 0 .2em}"
    ".byline-doc{color:#666;margin:0 0 .1em}.src{color:#999;font-size:.85em;margin:0 0 1.5em}"
    ".juan{font-size:1.1em;color:#8b6914;text-align:center;margin:1.5em 0 1em}"
    ".head{font-size:1.2em;color:#8b6914;text-align:center;margin:.5em 0 1em}"
    ".byline{color:#408080;text-align:right;margin:0 0 1.2em;font-size:.95em}"
    ".section{text-align:center;margin:1.5em 0 1em;font-size:1.05em}"
    ".prose{text-indent:2em;margin:0 0 1em}"
    ".verse{margin:0;padding-left:4em;text-indent:0}"
)


def _seg_html(seg: Segment) -> str:
    if seg.type == "break":
        return ""
    return f'<p class="{seg.type}">{escape(seg.text)}</p>'


def render_html(doc: ExportDoc) -> str:
    parts: list[str] = [
        "<!DOCTYPE html>",
        '<html lang="zh-Hant">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{escape(doc.title)}</title>",
        f"<style>{_HTML_STYLE}</style>",
        "</head>",
        "<body>",
        f"<h1>{escape(doc.title)}</h1>",
    ]
    byline = _byline(doc)
    if byline:
        parts.append(f'<p class="byline-doc">{escape(byline)}</p>')
    parts.append(f'<p class="src">CBETA {escape(doc.cbeta_id)}</p>')
    for j in doc.juans:
        parts.append(f'<h2 class="juan">{escape(_juan_label(j.juan_num))}</h2>')
        parts.extend(html for html in (_seg_html(s) for s in j.segments) if html)
    parts.append("</body>")
    parts.append("</html>")
    return "\n".join(parts) + "\n"


# --- DOCX --------------------------------------------------------------------

def render_docx(doc: ExportDoc) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = Document()
    d.add_heading(doc.title, level=0)
    byline = _byline(doc)
    if byline:
        d.add_paragraph(byline)
    d.add_paragraph(f"CBETA {doc.cbeta_id}")
    for j in doc.juans:
        d.add_heading(_juan_label(j.juan_num), level=1)
        for seg in j.segments:
            if seg.type == "break":
                continue
            if seg.type in ("head", "juan"):
                h = d.add_heading(seg.text, level=2)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            p = d.add_paragraph(seg.text)
            if seg.type == "section":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif seg.type == "byline":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif seg.type == "prose":
                p.paragraph_format.first_line_indent = Pt(24)
            elif seg.type == "verse":
                p.paragraph_format.left_indent = Pt(48)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


# --- EPUB --------------------------------------------------------------------

# Inline CSS shared by every juan document — same typography as the HTML export.
_EPUB_CSS = (
    "body{font-family:'Noto Serif SC',serif;line-height:1.8}"
    ".juan{font-size:1.1em;text-align:center;margin:1.5em 0 1em}"
    ".head{font-size:1.2em;text-align:center;margin:.5em 0 1em}"
    ".byline{text-align:right;margin:0 0 1.2em;font-size:.95em}"
    ".section{text-align:center;margin:1.5em 0 1em}"
    ".prose{text-indent:2em;margin:0 0 1em}"
    ".verse{margin:0;padding-left:4em}"
)


def _seg_xhtml(seg: Segment) -> str:
    if seg.type == "break":
        return ""
    return f'    <p class="{seg.type}">{xml_escape(seg.text)}</p>'


def _epub_juan_xhtml(title: str, j: ExportJuan) -> str:
    body = "\n".join(x for x in (_seg_xhtml(s) for s in j.segments) if x)
    return (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh-Hant">\n'
        f"  <head><title>{xml_escape(title)} {_juan_label(j.juan_num)}</title>\n"
        '  <meta charset="utf-8"/>\n'
        f"  <style>{_EPUB_CSS}</style></head>\n"
        "  <body>\n"
        f'    <h2 class="juan">{xml_escape(_juan_label(j.juan_num))}</h2>\n'
        f"{body}\n"
        "  </body>\n"
        "</html>\n"
    )


def render_epub(doc: ExportDoc) -> bytes:
    # Deterministic id so re-exports are byte-stable and the OPF is verifiable.
    book_id = f"urn:cbeta:{doc.cbeta_id}"
    juan_files = [(f"juan{j.juan_num}.xhtml", j) for j in doc.juans]

    manifest_items = "\n".join(
        f'    <item id="j{j.juan_num}" href="{fn}" media-type="application/xhtml+xml"/>'
        for fn, j in juan_files
    )
    spine_items = "\n".join(f'    <itemref idref="j{j.juan_num}"/>' for _, j in juan_files)
    byline = _byline(doc)
    creator = f"<dc:creator>{xml_escape(byline)}</dc:creator>" if byline else ""
    opf = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="bookid">\n'
        '  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">\n'
        f'    <dc:identifier id="bookid">{xml_escape(book_id)}</dc:identifier>\n'
        f"    <dc:title>{xml_escape(doc.title)}</dc:title>\n"
        "    <dc:language>zh-Hant</dc:language>\n"
        f"    {creator}\n"
        f"    <dc:source>CBETA {xml_escape(doc.cbeta_id)}</dc:source>\n"
        f'    <meta property="dcterms:modified">{_EPUB_MODIFIED}</meta>\n'
        "  </metadata>\n"
        "  <manifest>\n"
        '    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>\n'
        f"{manifest_items}\n"
        "  </manifest>\n"
        "  <spine>\n"
        f"{spine_items}\n"
        "  </spine>\n"
        "</package>\n"
    )

    nav_links = "\n".join(
        f'      <li><a href="{fn}">{xml_escape(_juan_label(j.juan_num))}</a></li>'
        for fn, j in juan_files
    )
    nav = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="zh-Hant">\n'
        f"  <head><title>{xml_escape(doc.title)}</title></head>\n"
        "  <body>\n"
        '    <nav epub:type="toc">\n'
        f"      <h1>{xml_escape(doc.title)}</h1>\n"
        "      <ol>\n"
        f"{nav_links}\n"
        "      </ol>\n"
        "    </nav>\n"
        "  </body>\n"
        "</html>\n"
    )

    container = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
        '  <rootfiles>\n'
        '    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>\n'
        "  </rootfiles>\n"
        "</container>\n"
    )

    buf = BytesIO()
    with ZipFile(buf, "w", ZIP_DEFLATED) as zf:
        # OCF rule: 'mimetype' first, stored uncompressed.
        zf.writestr(ZipInfo("mimetype"), "application/epub+zip", compress_type=ZIP_STORED)
        zf.writestr("META-INF/container.xml", container)
        zf.writestr("OEBPS/content.opf", opf)
        zf.writestr("OEBPS/nav.xhtml", nav)
        for fn, j in juan_files:
            zf.writestr(f"OEBPS/{fn}", _epub_juan_xhtml(doc.title, j))
    return buf.getvalue()


# --- dispatch ----------------------------------------------------------------

def render(doc: ExportDoc, fmt: str) -> tuple[bytes, str]:
    """Render ``doc`` to ``fmt``; returns (payload_bytes, media_type)."""
    if fmt not in FORMATS:
        raise ValueError(f"unsupported export format: {fmt!r}")
    _, media_type = FORMATS[fmt]
    if fmt == "txt":
        return render_txt(doc).encode("utf-8"), media_type
    if fmt == "html":
        return render_html(doc).encode("utf-8"), media_type
    if fmt == "docx":
        return render_docx(doc), media_type
    return render_epub(doc), media_type


# --- DB assembly -------------------------------------------------------------

def _split_lines(content: str) -> list[str]:
    return [ln.strip() for ln in content.split("\n") if ln.strip()]


async def build_export_doc(
    session: AsyncSession, text_id: int, juan_num: int | None
) -> ExportDoc | None:
    """Assemble an ExportDoc from stored content. ``juan_num=None`` exports the
    whole text; otherwise just that one juan. Returns None if the text is
    unknown or has no content."""
    result = await session.execute(select(BuddhistText).where(BuddhistText.id == text_id))
    bt = result.scalar_one_or_none()
    if bt is None:
        return None

    query = (
        select(TextContent)
        .where(TextContent.text_id == text_id, TextContent.lang == bt.lang)
        .order_by(TextContent.juan_num)
    )
    if juan_num is not None:
        query = query.where(TextContent.juan_num == juan_num)
    result = await session.execute(query)
    rows = result.scalars().all()
    if not rows:
        return None

    juans = [
        ExportJuan(
            juan_num=r.juan_num,
            lines=_split_lines(r.content or ""),
            segments=reflow(r.content or ""),
        )
        for r in rows
    ]
    return ExportDoc(
        title=bt.title_zh,
        cbeta_id=bt.cbeta_id,
        translator=bt.translator,
        dynasty=bt.dynasty,
        juans=juans,
    )
